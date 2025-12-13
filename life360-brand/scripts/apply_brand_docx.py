import argparse
import base64
import io
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuration Paths ---
SKILL_ROOT = Path(__file__).parent.parent
CONFIG_PATH = SKILL_ROOT / 'config.json'
LOGO_PATH = SKILL_ROOT / 'assets' / 'logo.txt'

def load_config():
    with open(CONFIG_PATH, 'r') as f:
        return json.load(f)

def load_logo():
    with open(LOGO_PATH, 'r') as f:
        return f.read().strip()

def apply_branding(input_path, output_path):
    print(f"Branding document: {input_path}")
    
    config = load_config()
    logo_b64 = load_logo()
    
    # Parse Colors
    primary_rgb = config['colors']['primary']['rgb']
    heading_rgb = config['colors']['text']['heading']['rgb']
    body_rgb = config['colors']['text']['body']['rgb']
    
    primary_color = RGBColor(*primary_rgb)
    heading_color = RGBColor(*heading_rgb)
    body_color = RGBColor(*body_rgb)
    
    brand_font = config['fonts']['primary']
    fallback_font = config['fonts']['fallback']

    try:
        doc = Document(input_path)
        
        # --- Add Header with Logo ---
        header = doc.sections[0].header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Decode Image
        logo_data = base64.b64decode(logo_b64)
        image_stream = io.BytesIO(logo_data)
        
        # Clear existing runs in header paragraph to avoid duplicates if re-running
        for run in paragraph.runs:
            run.text = ""
            
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Pt(100))

        # --- Apply Fonts and Colors ---
        for para in doc.paragraphs:
            style_name = para.style.name if para.style and para.style.name else ''
            is_heading = style_name.startswith('Heading')

            for run in para.runs:
                try:
                    run.font.name = brand_font
                except KeyError:
                    run.font.name = fallback_font

                if style_name == 'Title':
                    run.font.color.rgb = primary_color
                    run.font.size = Pt(26)
                elif is_heading:
                    run.font.color.rgb = heading_color
                else:
                    run.font.color.rgb = body_color
        
        doc.save(output_path)
        print(f"Successfully saved branded document to: {output_path}")
        
    except Exception as e:
        print(f"Error applying branding: {e}")
        raise e

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Apply Life360 Branding to a Word Document')
    parser.add_argument('input', help='Path to input .docx file')
    parser.add_argument('output', help='Path to save branded .docx file')
    
    args = parser.parse_args()
    apply_branding(args.input, args.output)